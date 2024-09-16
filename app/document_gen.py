from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def center_text_in_cell(cell):
    # Center text horizontally
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def replace_text_in_runs(paragraph, old_text, new_text):
    # Iterate over each run in the paragraph
    for run in paragraph.runs:
        if old_text in run.text:
            # Replace text in the run while preserving formatting
            run.text = run.text.replace(old_text, new_text)

def add_table_entires(table, entries):
    for entry in entries:
        row = table.add_row()  # Add a new row
        for i, cell_text in enumerate(entry):
            cell = row.cells[i]
            cell.text = cell_text

            # Center text in the cell
            center_text_in_cell(cell)

def replace_text_in_docx(input_file, output_file, entries, table_entries):
    # Load the document
    doc = Document(input_file)
    
    # Loop through each paragraph in the document
    for entry in entries.keys():
        old_text = entry
        new_text = entries[entry]
        for paragraph in doc.paragraphs:
            replace_text_in_runs(paragraph, old_text, new_text)
    
    table = doc.tables[0]
    add_table_entires(table, table_entries)

    # Save the modified document as a new file
    doc.save(output_file)




# Test the function
input_file = 'template_quotation.docx'
output_file = 'new_document.docx'
entries = {
    "[name]": "John Doe",
    "[company]": "ABC Inc.",
    "[date]": "01/01/2022",
    "[address]": "123 Main St",
    "[contact]": "123-456-7890",
    "[ref_no]": "Q12345",
    "[so_no]": "NIL",
    "[terms]": "30 days",
    "[LEW Name]": "Jane Smith",
    "[LEW No]": "12345",
    "[LEW Email]": "hello@hello.com",
    "[LEW Contact]": "987-654-3210",
}
table_entiries = [["1", "LEW Services", "1", "1000", "1000"]]

# Not used in production
if __name__ == '__main__':
    replace_text_in_docx(input_file, output_file, entries, table_entiries)

